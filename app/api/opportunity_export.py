import io
from typing import Any, Dict

from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import StreamingResponse

from app.auth.session import get_current_user_email
from app.api.auth_helpers import require_user_with_team, ensure_user_can_access_opportunity

router = APIRouter(prefix="/api/opportunities", tags=["opportunity-export"])


def _ascii_safe(text: str) -> str:
    """
    FPDF built-in fonts are latin-1 only; normalize/strip smart quotes and unicode chars.
    """
    if not text:
        return ""
    replacements = {
        "“": '"',
        "”": '"',
        "‘": "'",
        "’": "'",
        "–": "-",
        "—": "-",
        "•": "-",
        "…": "...",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    try:
        return text.encode("latin-1", "ignore").decode("latin-1")
    except Exception:
        return text


def _build_docx(payload: Dict[str, Any]) -> bytes:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail="python-docx is not installed") from exc

    cover = payload.get("cover_letter") or ""
    soq = payload.get("soq_body") or ""
    title = payload.get("title") or "Statement of Qualifications"
    agency = payload.get("agency") or ""

    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass

    doc.add_heading(title, level=1)
    if agency:
        doc.add_paragraph(agency)

    doc.add_heading("Cover Letter", level=2)
    for line in str(cover).splitlines():
        doc.add_paragraph(line)

    doc.add_heading("Statement of Qualifications", level=2)
    for line in str(soq).splitlines():
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_pdf(payload: Dict[str, Any]) -> bytes:
    try:
        from fpdf import FPDF  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail="fpdf2 is not installed") from exc

    cover = payload.get("cover_letter") or ""
    soq = payload.get("soq_body") or ""
    title = payload.get("title") or "Statement of Qualifications"
    agency = payload.get("agency") or ""

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, _ascii_safe(title), ln=1)
    if agency:
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 8, _ascii_safe(agency))

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 10, "Cover Letter", ln=1)
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, _ascii_safe(str(cover)))

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 10, "Statement of Qualifications", ln=1)
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, _ascii_safe(str(soq)))

    return pdf.output(dest="S").encode("latin-1", errors="replace")


@router.post("/{opportunity_id}/export")
async def export_documents(
    opportunity_id: str,
    payload: Dict[str, Any],
    format: str = Query("docx", regex="^(docx|pdf)$"),
    user=Depends(require_user_with_team),
):
    await ensure_user_can_access_opportunity(user, opportunity_id)
    if not payload:
        raise HTTPException(status_code=400, detail="Missing payload")
    if format == "docx":
        data = _build_docx(payload)
        filename = f"soq_{opportunity_id}.docx"
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    data = _build_pdf(payload)
    filename = f"soq_{opportunity_id}.pdf"
    return StreamingResponse(
        io.BytesIO(data),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
